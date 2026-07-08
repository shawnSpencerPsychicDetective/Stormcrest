from pathlib import Path

from .compress import compress


def universal_reader(absolute_folder_path, filename):
    """
    Read text from TXT, MD, PDF, DOCX, and ODT files.

    Args:
        absolute_folder_path (str): Absolute path to the folder.
        filename (str): File name including extension.

    Returns:
        str: Extracted text content.

    Raises:
        FileNotFoundError: If file does not exist.
        ValueError: If file type is unsupported.
    """

    file_path = Path(absolute_folder_path) / filename

    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    suffix = file_path.suffix.lower()

    # ------------------------
    # Plain Text Files
    # ------------------------
    if suffix in [".txt", ".md"]:
        return file_path.read_text(encoding="utf-8", errors="ignore")

    # ------------------------
    # PDF Files
    # ------------------------
    elif suffix == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ImportError("pypdf is required. Install with: pip install pypdf")

        reader = PdfReader(str(file_path))

        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)

        return "\n\n".join(pages)

    # ------------------------
    # DOCX Files
    # ------------------------
    elif suffix == ".docx":
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "python-docx is required. Install with: pip install python-docx"
            )

        doc = Document(str(file_path))

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Include table contents
        for table in doc.tables:
            for row in table.rows:
                paragraphs.append("\t".join(cell.text for cell in row.cells))

        return "\n".join(paragraphs)

    # ------------------------
    # ODT Files
    # ------------------------
    elif suffix == ".odt":
        try:
            from odf.opendocument import load
            from odf.text import P
        except ImportError:
            raise ImportError("odfpy is required. Install with: pip install odfpy")

        doc = load(str(file_path))

        paragraphs = []
        for para in doc.getElementsByType(P):
            text = "".join(
                node.data for node in para.childNodes if hasattr(node, "data")
            )
            if text.strip():
                paragraphs.append(text)

        return "\n".join(paragraphs)

    # ------------------------
    # Unsupported
    # ------------------------
    else:
        raise ValueError(
            f"Unsupported file type: {suffix}. Supported: .txt, .md, .pdf, .docx, .odt"
        )


def read(absolute_folder_path, filename):
    return compress(universal_reader(absolute_folder_path, filename))
